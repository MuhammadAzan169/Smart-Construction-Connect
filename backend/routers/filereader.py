"""
file_processor.py
Integrated file processing for AI Planning System - Web Service Only
=====================================================================
Professional, reusable file processor supporting multiple formats with OCR.
Designed for clean code, extensibility, and production use.

Features:
- Web Service Integration for FastAPI
- Multiple format support (PDF, Word, Excel, CSV, JSON, Images with OCR)
- Uses EasyOCR for image processing (no Tesseract dependency)
- Structured content extraction
- Error handling and logging
- Universal file upload with code file conversion to .txt

Usage: 
  Web Integration: Use FileProcessingService class
"""

import os
import sys
import json
import logging
import mimetypes
import pandas as pd
import tempfile
import warnings
import hashlib
import threading
import shutil
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple, Union
from datetime import datetime
from dataclasses import dataclass, asdict
from abc import ABC, abstractmethod
from cachetools import TTLCache

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# ==================== CONFIGURATION ====================
# Disable magic library to avoid Windows issues
MAGIC_AVAILABLE = False

# Code file extensions (must not be parsed/analyzed)
CODE_EXTENSIONS = {
    '.py', '.js', '.ts', '.jsx', '.tsx',
    '.java', '.cpp', '.c', '.h',
    '.cs', '.go', '.rs',
    '.php', '.sh', '.sql',
    '.yaml', '.yml'
}

# Set up logging early
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Conditional imports with fallbacks
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    try:
        import pdfplumber
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import easyocr
    EASYOCR_AVAILABLE = True
except (ImportError, KeyboardInterrupt, Exception) as e:
    EASYOCR_AVAILABLE = False
    logger.warning(f"EasyOCR not available: {e}")

# Check if file processor is available
FILE_PROCESSOR_AVAILABLE = all([
    PDF_AVAILABLE or not any(['pdf' in ext for ext in ['.pdf']]),  # PDF optional
    DOCX_AVAILABLE or not any(['doc' in ext for ext in ['.doc', '.docx']]),  # Word optional
    EXCEL_AVAILABLE or not any(['xls' in ext for ext in ['.xls', '.xlsx', '.csv']]),  # Excel optional
    EASYOCR_AVAILABLE or not any(['image' in ext for ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif', '.webp']])  # OCR optional
])

# ==================== DATA MODELS ====================

@dataclass
class FileMetadata:
    """Metadata for processed files"""
    filename: str
    file_type: str
    mime_type: str
    size_bytes: int
    upload_timestamp: str
    processing_time_ms: int
    page_count: Optional[int] = None
    language: Optional[str] = None
    dimensions: Optional[Tuple[int, int]] = None
    has_tables: bool = False
    error: Optional[str] = None

@dataclass
class ExtractionResult:
    """Result of file extraction"""
    success: bool
    content: str
    structured_data: Optional[Dict[str, Any]] = None
    metadata: Optional[FileMetadata] = None
    tables: Optional[List[Dict[str, Any]]] = None
    keywords: Optional[List[str]] = None
    summary: Optional[str] = None
    error_message: Optional[str] = None

class FileProcessorConfig:
    """Configuration for file processing"""
    
    def __init__(
        self,
        use_ocr: bool = True,
        ocr_languages: List[str] = None,
        extract_tables: bool = True,
        extract_keywords: bool = True,
        generate_summary: bool = False,
        max_file_size_mb: int = int(os.getenv('MAX_FILE_SIZE_MB', 5)),
        output_format: str = "text",  # "text", "json", "both"
        cleanup_temp_files: bool = True,
        verbose_logging: bool = False
    ):
        self.use_ocr = use_ocr
        # Clean language codes - remove quotes and whitespace
        if ocr_languages:
            self.ocr_languages = [lang.strip().strip("'\"") for lang in ocr_languages]
        else:
            self.ocr_languages = ['en']
        self.extract_tables = extract_tables
        self.extract_keywords = extract_keywords
        self.generate_summary = generate_summary
        self.max_file_size_mb = max_file_size_mb
        self.output_format = output_format
        self.cleanup_temp_files = cleanup_temp_files
        self.verbose_logging = verbose_logging

# ==================== ABSTRACT BASE CLASS ====================

class BaseFileHandler(ABC):
    """Abstract base class for file handlers"""
    
    def __init__(self, config: FileProcessorConfig):
        self.config = config
        self.logger = logging.getLogger(self.__class__.__name__)
    
    @abstractmethod
    def can_handle(self, file_path: Path, mime_type: str) -> bool:
        """Check if handler can process this file type"""
        pass
    
    @abstractmethod
    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract content from file"""
        pass
    
    def _create_metadata(self, file_path: Path, mime_type: str) -> FileMetadata:
        """Create metadata object"""
        stat = file_path.stat()
        return FileMetadata(
            filename=file_path.name,
            file_type=file_path.suffix.lower(),
            mime_type=mime_type,
            size_bytes=stat.st_size,
            upload_timestamp=datetime.now().isoformat(),
            processing_time_ms=0  # Will be updated after processing
        )
    
    def _extract_keywords(self, text: str, max_keywords: int = 10) -> List[str]:
        """Extract keywords from text using simple frequency analysis"""
        if not text or len(text.split()) < 5:
            return []
        
        # Remove common stopwords
        stopwords = set([
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
            'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'should',
            'can', 'could', 'may', 'might', 'must', 'shall'
        ])
        
        # Tokenize and count
        words = text.lower().split()
        word_count = {}
        
        for word in words:
            # Clean word
            word = ''.join(c for c in word if c.isalnum())
            if len(word) > 2 and word not in stopwords:
                word_count[word] = word_count.get(word, 0) + 1
        
        # Get top keywords
        sorted_words = sorted(word_count.items(), key=lambda x: x[1], reverse=True)
        return [word for word, count in sorted_words[:max_keywords]]
    
    def _generate_summary(self, text: str, max_sentences: int = 3) -> str:
        """Generate a simple summary of the text"""
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        
        if len(sentences) <= max_sentences:
            return '. '.join(sentences) + ('.' if sentences else '')
        
        # Simple algorithm: take first and last sentences
        summary_sentences = sentences[:max_sentences]
        return '. '.join(summary_sentences) + '.'

# ==================== CONCRETE HANDLERS ====================

class PDFHandler(BaseFileHandler):
    """Handler for PDF files"""
    
    def __init__(self, config: FileProcessorConfig):
        super().__init__(config)
        if not PDF_AVAILABLE:
            self.logger.warning("PDF processing libraries not available. Install PyMuPDF or pdfplumber.")
    
    def can_handle(self, file_path: Path, mime_type: str) -> bool:
        return mime_type == 'application/pdf' or file_path.suffix.lower() in ['.pdf']
    
    def extract(self, file_path: Path) -> ExtractionResult:
        start_time = datetime.now()
        metadata = self._create_metadata(file_path, 'application/pdf')
        
        try:
            content_parts = []
            tables = []
            
            # Try PyMuPDF first, fallback to pdfplumber
            if PDF_AVAILABLE:
                if 'fitz' in sys.modules:
                    content, extracted_tables = self._extract_with_pymupdf(file_path)
                else:
                    content, extracted_tables = self._extract_with_pdfplumber(file_path)
                
                content_parts.append(content)
                if extracted_tables and self.config.extract_tables:
                    tables.extend(extracted_tables)
                
                # Extract text
                full_text = '\n\n'.join(content_parts)
                
                # Update metadata
                metadata.page_count = len(content_parts)
                metadata.has_tables = len(tables) > 0
                
                # Generate keywords if requested
                keywords = self._extract_keywords(full_text) if self.config.extract_keywords else None
                
                # Generate summary if requested
                summary = self._generate_summary(full_text) if self.config.generate_summary else None
                
                processing_time = int((datetime.now() - start_time).total_seconds() * 1000)
                metadata.processing_time_ms = processing_time
                
                return ExtractionResult(
                    success=True,
                    content=full_text,
                    structured_data={"pages": content_parts, "tables": tables},
                    metadata=metadata,
                    tables=tables,
                    keywords=keywords,
                    summary=summary
                )
            else:
                return ExtractionResult(
                    success=False,
                    content="",
                    metadata=metadata,
                    error_message="PDF processing libraries not available. Install PyMuPDF or pdfplumber."
                )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                metadata=metadata,
                error_message=f"PDF extraction error: {str(e)}"
            )
    
    def _extract_with_pymupdf(self, file_path: Path) -> Tuple[str, List[Dict]]:
        """Extract using PyMuPDF"""
        import fitz
        doc = fitz.open(file_path)
        pages_text = []
        tables = []
        
        for page_num, page in enumerate(doc):
            text = page.get_text()
            pages_text.append(f"=== Page {page_num + 1} ===\n{text}")
            
            # Extract tables if requested
            if self.config.extract_tables:
                tabs = page.find_tables()
                for tab in tabs:
                    table_data = tab.extract()
                    if table_data:
                        tables.append({
                            "page": page_num + 1,
                            "data": table_data,
                            "headers": table_data[0] if table_data else []
                        })
        
        doc.close()
        return '\n'.join(pages_text), tables
    
    def _extract_with_pdfplumber(self, file_path: Path) -> Tuple[str, List[Dict]]:
        """Extract using pdfplumber"""
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            pages_text = []
            tables = []
            
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                pages_text.append(f"=== Page {page_num + 1} ===\n{text}")
                
                # Extract tables if requested
                if self.config.extract_tables:
                    page_tables = page.extract_tables()
                    for table_num, table in enumerate(page_tables):
                        tables.append({
                            "page": page_num + 1,
                            "table": table_num + 1,
                            "data": table,
                            "headers": table[0] if table else []
                        })
            
            return '\n'.join(pages_text), tables

class WordHandler(BaseFileHandler):
    """Handler for Word documents"""
    
    def __init__(self, config: FileProcessorConfig):
        super().__init__(config)
        if not DOCX_AVAILABLE:
            self.logger.warning("python-docx not available. Word processing disabled.")
    
    def can_handle(self, file_path: Path, mime_type: str) -> bool:
        return (mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'] 
                or file_path.suffix.lower() in ['.doc', '.docx'])
    
    def extract(self, file_path: Path) -> ExtractionResult:
        start_time = datetime.now()
        metadata = self._create_metadata(file_path, 
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if file_path.suffix == '.docx' 
            else 'application/msword')
        
        if not DOCX_AVAILABLE:
            return ExtractionResult(
                success=False,
                content="",
                metadata=metadata,
                error_message="Word processing library not available. Install python-docx."
            )
        
        try:
            from docx import Document
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            tables = []
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        "table": table_num + 1,
                        "data": table_data,
                        "headers": table_data[0] if table_data else []
                    })
            
            full_text = '\n\n'.join(paragraphs)
            metadata.has_tables = len(tables) > 0
            
            # Generate keywords and summary
            keywords = self._extract_keywords(full_text) if self.config.extract_keywords else None
            summary = self._generate_summary(full_text) if self.config.generate_summary else None
            
            processing_time = int((datetime.now() - start_time).total_seconds() * 1000)
            metadata.processing_time_ms = processing_time
            
            return ExtractionResult(
                success=True,
                content=full_text,
                structured_data={"paragraphs": paragraphs, "tables": tables},
                metadata=metadata,
                tables=tables,
                keywords=keywords,
                summary=summary
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                metadata=metadata,
                error_message=f"Word document extraction error: {str(e)}"
            )

class ExcelHandler(BaseFileHandler):
    """Handler for Excel and CSV files"""
    
    def can_handle(self, file_path: Path, mime_type: str) -> bool:
        excel_mimes = [
            'application/vnd.ms-excel',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'text/csv'
        ]
        excel_exts = ['.xls', '.xlsx', '.csv']
        return (mime_type in excel_mimes or 
                file_path.suffix.lower() in excel_exts)
    
    def extract(self, file_path: Path) -> ExtractionResult:
        start_time = datetime.now()
        suffix = file_path.suffix.lower()
        
        if suffix == '.csv':
            mime_type = 'text/csv'
        elif suffix == '.xls':
            mime_type = 'application/vnd.ms-excel'
        else:
            mime_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        
        metadata = self._create_metadata(file_path, mime_type)
        
        try:
            if suffix == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path, engine='openpyxl' if EXCEL_AVAILABLE else None)
            
            # Extract content as text
            content = self._dataframe_to_text(df)
            
            # Convert to structured data
            structured_data = {
                "columns": df.columns.tolist(),
                "data": df.fillna('').to_dict(orient='records'),
                "shape": list(df.shape),
                "dtypes": df.dtypes.astype(str).to_dict()
            }
            
            tables = [{
                "sheet": "Sheet1",
                "data": df.fillna('').values.tolist(),
                "headers": df.columns.tolist()
            }]
            
            metadata.has_tables = True
            
            # Extract keywords from column names and sample data
            sample_text = ' '.join(df.columns.tolist()) + ' ' + ' '.join(
                df.head(10).fillna('').astype(str).values.flatten().tolist()
            )
            keywords = self._extract_keywords(sample_text) if self.config.extract_keywords else None
            
            processing_time = int((datetime.now() - start_time).total_seconds() * 1000)
            metadata.processing_time_ms = processing_time
            
            return ExtractionResult(
                success=True,
                content=content,
                structured_data=structured_data,
                metadata=metadata,
                tables=tables,
                keywords=keywords
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                metadata=metadata,
                error_message=f"Excel/CSV extraction error: {str(e)}"
            )
    
    def _dataframe_to_text(self, df: pd.DataFrame) -> str:
        """Convert DataFrame to readable text"""
        lines = []
        
        # Add headers
        lines.append(" | ".join(str(col) for col in df.columns))
        lines.append("-" * (sum(len(str(col)) for col in df.columns) + len(df.columns) * 3))
        
        # Add rows (limit to 100 for performance)
        for _, row in df.head(100).iterrows():
            lines.append(" | ".join(str(val) for val in row))
        
        if len(df) > 100:
            lines.append(f"\n... and {len(df) - 100} more rows")
        
        return '\n'.join(lines)

class JSONHandler(BaseFileHandler):
    """Handler for JSON files"""
    
    def can_handle(self, file_path: Path, mime_type: str) -> bool:
        return (mime_type == 'application/json' or 
                file_path.suffix.lower() == '.json')
    
    def extract(self, file_path: Path) -> ExtractionResult:
        start_time = datetime.now()
        metadata = self._create_metadata(file_path, 'application/json')
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert to formatted string
            content = json.dumps(data, indent=2, ensure_ascii=False)
            
            # Extract text content from JSON values
            text_content = self._extract_text_from_json(data)
            
            # Generate keywords
            keywords = self._extract_keywords(text_content) if self.config.extract_keywords else None
            
            processing_time = int((datetime.now() - start_time).total_seconds() * 1000)
            metadata.processing_time_ms = processing_time
            
            return ExtractionResult(
                success=True,
                content=content,
                structured_data=data,
                metadata=metadata,
                keywords=keywords
            )
            
        except json.JSONDecodeError as e:
            return ExtractionResult(
                success=False,
                content="",
                metadata=metadata,
                error_message=f"Invalid JSON: {str(e)}"
            )
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                metadata=metadata,
                error_message=f"JSON extraction error: {str(e)}"
            )
    
    def _extract_text_from_json(self, data: Any) -> str:
        """Extract all text values from JSON structure"""
        if isinstance(data, dict):
            texts = []
            for key, value in data.items():
                texts.append(str(key))
                texts.extend(self._extract_text_from_json(value))
            return ' '.join(texts)
        elif isinstance(data, list):
            texts = []
            for item in data:
                texts.extend(self._extract_text_from_json(item))
            return ' '.join(texts)
        elif isinstance(data, str):
            return data
        else:
            return str(data)

class ImageHandler(BaseFileHandler):
    """Handler for images with EasyOCR"""
    
    def __init__(self, config: FileProcessorConfig):
        super().__init__(config)
        self.ocr_reader = None
        self._initialize_ocr()
    
    def _initialize_ocr(self):
        """Initialize EasyOCR reader"""
        if not self.config.use_ocr:
            return
        
        if EASYOCR_AVAILABLE:
            try:
                # Ensure language codes are clean
                clean_languages = [lang.strip().strip("'\"") for lang in self.config.ocr_languages]
                import easyocr
                self.ocr_reader = easyocr.Reader(clean_languages)
                self.logger.info(f"Initialized EasyOCR with languages: {clean_languages}")
            except Exception as e:
                self.logger.warning(f"Failed to initialize EasyOCR: {e}")
                self.ocr_reader = None
        else:
            self.logger.warning("EasyOCR not available. Install with: pip install easyocr")
    
    def can_handle(self, file_path: Path, mime_type: str) -> bool:
        # Updated to include WebP format
        image_mimes = ['image/png', 'image/jpeg', 'image/jpg', 'image/tiff', 'image/bmp', 'image/webp']
        image_exts = ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif', '.webp']
        return (mime_type in image_mimes or 
                file_path.suffix.lower() in image_exts)
    
    def extract(self, file_path: Path) -> ExtractionResult:
        start_time = datetime.now()
        
        # Determine correct MIME type
        if file_path.suffix.lower() == '.webp':
            mime_type = 'image/webp'
        elif file_path.suffix.lower() in ['.jpg', '.jpeg']:
            mime_type = 'image/jpeg'
        else:
            mime_type = 'image/png'
            
        metadata = self._create_metadata(file_path, mime_type)
        
        if not self.config.use_ocr:
            return ExtractionResult(
                success=False,
                content="",
                metadata=metadata,
                error_message="OCR is disabled in configuration"
            )
        
        if not EASYOCR_AVAILABLE:
            return ExtractionResult(
                success=False,
                content="",
                metadata=metadata,
                error_message="EasyOCR not available. Install with: pip install easyocr"
            )
        
        try:
            # Extract text with EasyOCR
            if self.ocr_reader:
                result = self._extract_with_easyocr(file_path)
            else:
                return ExtractionResult(
                    success=False,
                    content="",
                    metadata=metadata,
                    error_message="EasyOCR reader not initialized"
                )
            
            if not result.strip():
                metadata.error = "No text detected in image"
            
            # Extract keywords
            keywords = self._extract_keywords(result) if self.config.extract_keywords else None
            
            # Generate summary
            summary = self._generate_summary(result) if self.config.generate_summary else None
            
            processing_time = int((datetime.now() - start_time).total_seconds() * 1000)
            metadata.processing_time_ms = processing_time
            
            return ExtractionResult(
                success=True,
                content=result,
                metadata=metadata,
                keywords=keywords,
                summary=summary
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                metadata=metadata,
                error_message=f"Image OCR error: {str(e)}"
            )
    
    def _extract_with_easyocr(self, file_path: Path) -> str:
        """Extract text using EasyOCR"""
        result = self.ocr_reader.readtext(str(file_path), detail=0)
        return '\n'.join(result)

class TextHandler(BaseFileHandler):
    """Handler for plain text files"""
    
    def can_handle(self, file_path: Path, mime_type: str) -> bool:
        return (mime_type.startswith('text/') or 
                file_path.suffix.lower() in ['.txt', '.md', '.rtf'])
    
    def extract(self, file_path: Path) -> ExtractionResult:
        start_time = datetime.now()
        metadata = self._create_metadata(file_path, 'text/plain')
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return ExtractionResult(
                    success=False,
                    content="",
                    metadata=metadata,
                    error_message="Could not decode file with any supported encoding"
                )
            
            # Clean and normalize text
            content = self._clean_text(content)
            
            # Extract keywords
            keywords = self._extract_keywords(content) if self.config.extract_keywords else None
            
            # Generate summary
            summary = self._generate_summary(content) if self.config.generate_summary else None
            
            processing_time = int((datetime.now() - start_time).total_seconds() * 1000)
            metadata.processing_time_ms = processing_time
            
            return ExtractionResult(
                success=True,
                content=content,
                metadata=metadata,
                keywords=keywords,
                summary=summary
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                metadata=metadata,
                error_message=f"Text file extraction error: {str(e)}"
            )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Remove empty lines
        cleaned_lines = []
        for line in lines:
            if line:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)

# ==================== MAIN PROCESSOR CLASS ====================

class ModularFileProcessor:
    """
    Main file processor that orchestrates all file handlers
    """
    
    def __init__(self, config: Optional[FileProcessorConfig] = None):
        self.config = config or FileProcessorConfig()
        self.logger = self._setup_logging()
        self.cache = TTLCache(maxsize=50, ttl=600)  # 10-minute cache
        self.processing_lock = threading.RLock()
        
        # Initialize handlers
        self.handlers = [
            PDFHandler(self.config),
            WordHandler(self.config),
            ExcelHandler(self.config),
            JSONHandler(self.config),
            ImageHandler(self.config),
            TextHandler(self.config)
        ]
    
    def _setup_logging(self) -> logging.Logger:
        """Setup logging configuration"""
        logger = logging.getLogger('FileProcessor')
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
            
            if self.config.verbose_logging:
                logger.setLevel(logging.DEBUG)
            else:
                logger.setLevel(logging.INFO)
        
        return logger
    
    def detect_file_type(self, file_path: Path) -> Tuple[str, str]:
        """
        Detect file type using mimetypes
        Returns: (mime_type, file_extension)
        """
        # Method 1: Use mimetypes
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        # Method 2: Fallback to extension if mimetypes fails
        if not mime_type:
            extension = file_path.suffix.lower()
            if extension in ['.pdf']:
                mime_type = 'application/pdf'
            elif extension in ['.doc', '.docx']:
                mime_type = 'application/msword'
            elif extension in ['.xls', '.xlsx']:
                mime_type = 'application/vnd.ms-excel'
            elif extension == '.csv':
                mime_type = 'text/csv'
            elif extension == '.json':
                mime_type = 'application/json'
            elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:  # Added .webp
                mime_type = 'image/jpeg' if extension in ['.jpg', '.jpeg'] else 'image/png' if extension in ['.png'] else 'image/webp'
            elif extension in ['.txt', '.md', '.rtf']:
                mime_type = 'text/plain'
            else:
                mime_type = 'application/octet-stream'
        
        return mime_type, file_path.suffix.lower()
    
    def validate_file(self, file_path: Path) -> Tuple[bool, Optional[str]]:
        """Validate file before processing"""
        if not file_path.exists():
            return False, f"File not found: {file_path}"
        
        if not file_path.is_file():
            return False, f"Path is not a file: {file_path}"
        
        # Check file size
        max_size = self.config.max_file_size_mb * 1024 * 1024
        if file_path.stat().st_size > max_size:
            return False, f"File too large. Maximum size is {self.config.max_file_size_mb}MB"
        
        return True, None
    
    def is_code_file(self, file_path: Path) -> bool:
        """Check if file is a code file that should only be converted to .txt"""
        return file_path.suffix.lower() in CODE_EXTENSIONS
    
    def _get_file_hash(self, file_path: Path) -> str:
        """Generate hash for file caching"""
        try:
            hasher = hashlib.md5()
            with open(file_path, 'rb') as f:
                # Read first 1MB for hash (faster than entire file)
                chunk = f.read(1024 * 1024)
                hasher.update(chunk)
                hasher.update(str(file_path.stat().st_size).encode())
                hasher.update(file_path.suffix.encode())
            return hasher.hexdigest()
        except:
            return str(file_path)
    
    def process_file(self, file_path: Union[str, Path]) -> ExtractionResult:
        """Process file with caching and streaming"""
        file_path = Path(file_path)
        
        # Generate cache key from file hash
        cache_key = self._get_file_hash(file_path)
        
        # Check cache
        with self.processing_lock:
            if cache_key in self.cache:
                self.logger.info(f"Cache hit for {file_path.name}")
                return self.cache[cache_key]
        
        # Validate
        is_valid, error_msg = self.validate_file(file_path)
        if not is_valid:
            return ExtractionResult(
                success=False,
                content="",
                error_message=error_msg
            )
        
        # Check if code file
        if self.is_code_file(file_path):
            return self.convert_code_to_txt(file_path)
        
        # Process with appropriate handler
        mime_type, extension = self.detect_file_type(file_path)
        
        # Find appropriate handler
        handler = None
        for h in self.handlers:
            if h.can_handle(file_path, mime_type):
                handler = h
                break
        
        if not handler:
            return ExtractionResult(
                success=False,
                content="",
                error_message=f"Unsupported file type: {mime_type}"
            )
        
        # Process file
        try:
            result = handler.extract(file_path)
            
            # Cache successful results
            if result.success:
                with self.processing_lock:
                    self.cache[cache_key] = result
            
            return result
            
        except Exception as e:
            self.logger.error(f"Processing error: {e}")
            return ExtractionResult(
                success=False,
                content="",
                error_message=f"Processing error: {str(e)}"
            )
    
    def convert_code_to_txt(self, file_path: Path) -> ExtractionResult:
        """Convert code file to txt with streaming"""
        start_time = datetime.now()
        
        # Create temp output file
        temp_dir = tempfile.mkdtemp()
        txt_path = Path(temp_dir) / f"{file_path.stem}.txt"
        
        try:
            # Stream file line by line
            lines_processed = 0
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as src:
                with open(txt_path, 'w', encoding='utf-8') as dst:
                    # Write header
                    dst.write(f"# Converted from: {file_path.name}\n")
                    dst.write(f"# Converted at: {datetime.now().isoformat()}\n")
                    dst.write(f"# Original size: {file_path.stat().st_size} bytes\n")
                    dst.write(f"# Lines processed: [counting]\n\n")
                    
                    # Process file in chunks
                    buffer = []
                    for line in src:
                        buffer.append(line)
                        lines_processed += 1
                        
                        # Write in chunks of 1000 lines
                        if len(buffer) >= 1000:
                            dst.writelines(buffer)
                            buffer = []
                    
                    # Write remaining
                    if buffer:
                        dst.writelines(buffer)
            
            # Update line count
            with open(txt_path, 'r') as f:
                content = f.readlines()
            
            # Update line count in header
            content[3] = f"# Lines processed: {lines_processed}\n\n"
            
            with open(txt_path, 'w') as f:
                f.writelines(content)
            
            # Read final content
            with open(txt_path, 'r') as f:
                final_content = f.read()
            
            processing_time = int((datetime.now() - start_time).total_seconds() * 1000)
            
            metadata = FileMetadata(
                filename=txt_path.name,
                file_type='.txt',
                mime_type='text/plain',
                size_bytes=txt_path.stat().st_size,
                upload_timestamp=datetime.now().isoformat(),
                processing_time_ms=processing_time,
                language=file_path.suffix.replace('.', '')
            )
            
            return ExtractionResult(
                success=True,
                content=final_content,
                metadata=metadata
            )
            
        except Exception as e:
            self.logger.error(f"Code conversion failed: {e}")
            return ExtractionResult(
                success=False,
                content="",
                error_message=f"Conversion failed: {str(e)}"
            )
        finally:
            # Cleanup temp directory
            try:
                shutil.rmtree(temp_dir)
            except:
                pass

# ==================== WEB SERVICE INTEGRATION ====================

class FileProcessingService:
    """Service for processing uploaded files in web applications"""
    
    def __init__(self):
        self.processor = None
        self.config = None
        
        if FILE_PROCESSOR_AVAILABLE:
            self.config = FileProcessorConfig(
                use_ocr=True,
                ocr_languages=['en'],
                extract_tables=True,
                extract_keywords=True,
                generate_summary=True,
                max_file_size_mb=int(os.getenv('MAX_FILE_SIZE_MB', 5)),
                output_format="text",
                verbose_logging=False
            )
            self.processor = ModularFileProcessor(self.config)
    
    async def process_uploaded_file(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """Process an uploaded file and return extracted content"""
        if not FILE_PROCESSOR_AVAILABLE or not self.processor:
            return {
                "success": False,
                "error": "File processor not available",
                "content": ""
            }
        
        try:
            # Save file to temporary location
            with tempfile.NamedTemporaryFile(suffix=Path(filename).suffix, delete=False) as tmp:
                tmp.write(file_data)
                tmp_path = tmp.name
            
            # Process the file
            result = self.processor.process_file(tmp_path)
            
            # Clean up temp file
            try:
                os.unlink(tmp_path)
            except:
                pass
            
            if result.success:
                # Format the result for display
                formatted_content = self._format_extracted_content(result, filename)
                
                return {
                    "success": True,
                    "content": formatted_content,
                    "raw_content": result.content[:5000],  # Limit for context
                    "metadata": {
                        "filename": filename,
                        "type": result.metadata.file_type if result.metadata else "unknown",
                        "size": result.metadata.size_bytes if result.metadata else 0,
                        "keywords": result.keywords or [],
                        "summary": result.summary or ""
                    }
                }
            else:
                return {
                    "success": False,
                    "error": result.error_message,
                    "content": ""
                }
                
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            return {
                "success": False,
                "error": str(e),
                "content": ""
            }
    
    def _format_extracted_content(self, result: ExtractionResult, filename: str) -> str:
        """Format extracted content for display"""
        lines = []
        
        # Header
        lines.append(f"📄 **File:** {filename}")
        
        if result.metadata:
            size_kb = result.metadata.size_bytes / 1024
            lines.append(f"📊 **Size:** {size_kb:.1f} KB | **Type:** {result.metadata.file_type}")
        
        if result.keywords:
            lines.append(f"🔑 **Key Topics:** {', '.join(result.keywords[:8])}")
        
        if result.summary:
            lines.append(f"\n📝 **Summary:** {result.summary}")
        
        # Content preview
        lines.append(f"\n📋 **Extracted Content Preview:**")
        
        # For code files, show conversion notice
        if result.metadata and result.metadata.language and result.metadata.file_type in ['.txt']:
            if result.metadata.language in ['py', 'js', 'java', 'cpp', 'c', 'cs', 'go', 'rs', 'php', 'sql', 'yaml', 'yml']:
                lines.append(f"*Note: Code file converted to text without analysis*")
        
        # Add content (truncated for display)
        preview = result.content[:1500]  # Limit preview
        if len(result.content) > 1500:
            preview += "\n\n... [content truncated for display] ..."
        
        lines.append(f"\n{preview}")
        
        return "\n".join(lines)
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return [
            "📄 Documents: .pdf, .doc, .docx",
            "📊 Spreadsheets: .xls, .xlsx, .csv",
            "🖼️ Images: .png, .jpg, .jpeg, .webp (with OCR)",
            "📝 Text: .txt, .md, .json",
            "💻 Code: .py, .js, .java, .cpp, .c, .cs, .go, .rs, .php, .sql, .yaml, .yml"
        ]

# Global instance for web service
file_processor = FileProcessingService()